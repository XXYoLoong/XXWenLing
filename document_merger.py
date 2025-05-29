import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from tkinter import ttk
import os
import sys
import json
from docx import Document
import win32com.client
import pythoncom
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def resource_path(relative_path):
    """ 获取资源的绝对路径 """
    try:
        # PyInstaller创建临时文件夹,将路径存储在_MEIPASS中
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

CONFIG_FILE = os.path.join(os.path.expanduser("~"), "document_merger_config.json")

PRIMARY_COLOR = "#2563eb"  # 明快蓝色
ACCENT_COLOR = "#ff9800"   # 橙色高亮
SUCCESS_COLOR = "#43a047"  # 绿色
ERROR_COLOR = "#e53935"    # 红色
BG_COLOR = "#FFFFFF"       # 纯白背景
LABEL_COLOR = "#222222"    # 文字色
TITLE_FONT = ("微软雅黑", 20, "bold")
LABEL_FONT = ("微软雅黑", 13)
BTN_FONT = ("微软雅黑", 13, "bold")

class CustomButton(ttk.Button):
    def __init__(self, master=None, **kw):
        super().__init__(master, **kw)
        self.default_bg = PRIMARY_COLOR
        self.hover_bg = ACCENT_COLOR
        self.configure(style="TButton")
        self.bind("<Enter>", self.on_enter)
        self.bind("<Leave>", self.on_leave)
    def on_enter(self, e):
        self.configure(style="Hover.TButton")
    def on_leave(self, e):
        self.configure(style="TButton")

class DocumentMerger:
    def __init__(self, root):
        self.root = root
        self.root.title("文档合并工具")
        self.root.geometry("700x560")
        self.root.configure(bg=BG_COLOR)
        self.root.resizable(False, False)
        self.center_window()
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)
        self.load_config()
        self.show_instruction_dialog()  # 每次启动都弹窗
        
        # 设置窗口图标（如果有的话）
        # self.root.iconbitmap(resource_path("icon.ico"))
        
        style = ttk.Style()
        style.theme_use('clam')
        style.configure("TButton",
                        font=BTN_FONT,
                        background=PRIMARY_COLOR,
                        foreground="#fff",
                        borderwidth=0,
                        focusthickness=0,
                        focuscolor=PRIMARY_COLOR,
                        padding=14)
        style.map("TButton",
                  background=[('active', ACCENT_COLOR), ('!active', PRIMARY_COLOR)])
        style.configure("Hover.TButton",
                        font=BTN_FONT,
                        background=ACCENT_COLOR,
                        foreground="#fff",
                        borderwidth=0)
        style.configure("Custom.TLabel",
                        font=LABEL_FONT,
                        background=BG_COLOR,
                        foreground=LABEL_COLOR)
        style.configure("Title.TLabel",
                        font=TITLE_FONT,
                        background=BG_COLOR,
                        foreground=PRIMARY_COLOR)
        style.configure("Status.TLabel",
                        font=LABEL_FONT,
                        background=BG_COLOR,
                        foreground=ACCENT_COLOR)
        style.configure("Success.TLabel",
                        font=LABEL_FONT,
                        background=BG_COLOR,
                        foreground=SUCCESS_COLOR)
        style.configure("Error.TLabel",
                        font=LABEL_FONT,
                        background=BG_COLOR,
                        foreground=ERROR_COLOR)
        style.configure("TProgressbar",
                        troughcolor="#E0E3E9",
                        bordercolor=BG_COLOR,
                        background=PRIMARY_COLOR,
                        lightcolor=PRIMARY_COLOR,
                        darkcolor=PRIMARY_COLOR)

        self.main_frame = ttk.Frame(root, padding="30 30 30 30", style="Custom.TLabel")
        self.main_frame.place(relx=0.5, rely=0.5, anchor="center")

        ttk.Label(self.main_frame, text="文档合并助手", style="Title.TLabel").grid(row=0, column=0, columnspan=2, pady=(0, 28))

        self.help_btn = CustomButton(
            self.main_frame,
            text="查看使用方法",
            command=self.safe_call(self.show_instruction_dialog)
        )
        self.help_btn.grid(row=1, column=0, pady=8, sticky="ew")
        
        self.reset_btn = CustomButton(
            self.main_frame,
            text="搞错了，再来！",
            command=self.safe_call(self.reset_all)
        )
        self.reset_btn.grid(row=1, column=1, pady=8, padx=16, sticky="ew")
        
        self.select_dir_btn = CustomButton(
            self.main_frame, 
            text="选择文档目录",
            command=self.safe_call(self.select_directory)
        )
        self.select_dir_btn.grid(row=2, column=0, pady=16, sticky="ew")
        
        self.dir_label = ttk.Label(self.main_frame, text="未选择目录", style="Custom.TLabel")
        self.dir_label.grid(row=2, column=1, pady=16, sticky="w")
        
        self.name_btn = CustomButton(
            self.main_frame,
            text="命名合并文档",
            command=self.safe_call(self.set_output_name),
            state="disabled"
        )
        self.name_btn.grid(row=3, column=0, pady=16, sticky="ew")
        
        self.save_dir_btn = CustomButton(
            self.main_frame,
            text="选择保存的目标目录",
            command=self.safe_call(self.select_save_directory),
            state="disabled"
        )
        self.save_dir_btn.grid(row=3, column=1, pady=16, padx=16, sticky="ew")
        
        self.merge_btn = CustomButton(
            self.main_frame,
            text="开始合成",
            command=self.safe_call(self.merge_documents),
            state="disabled"
        )
        self.merge_btn.grid(row=4, column=0, pady=18, sticky="ew")
        
        self.progress = ttk.Progressbar(
            self.main_frame,
            orient="horizontal",
            length=360,
            mode="determinate",
            style="TProgressbar"
        )
        self.progress.grid(row=5, column=0, columnspan=2, pady=18, sticky="ew")
        
        self.status_label = ttk.Label(self.main_frame, text="", style="Status.TLabel")
        self.status_label.grid(row=6, column=0, columnspan=2, pady=8)
        
        version_label = ttk.Label(self.main_frame, text="版本 1.0.0", style="Custom.TLabel")
        version_label.grid(row=7, column=0, columnspan=2, pady=8)
        
        copyright_label = ttk.Label(
            self.main_frame, 
            text="© 倪家诚 2025/5/29",
            style="Custom.TLabel",
            font=("微软雅黑", 11)
        )
        copyright_label.grid(row=8, column=0, columnspan=2, pady=8)
        
        self.selected_dir = None
        self.output_name = None
        self.save_dir = None
        self.supported_extensions = ['.doc', '.docx']

    def center_window(self):
        self.root.update_idletasks()
        w = self.root.winfo_width()
        h = self.root.winfo_height()
        ws = self.root.winfo_screenwidth()
        hs = self.root.winfo_screenheight()
        x = (ws // 2) - (w // 2)
        y = (hs // 2) - (h // 2)
        self.root.geometry(f'{w}x{h}+{x}+{y}')

    def load_config(self):
        self.config = {"show_instruction": True}
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, "r", encoding="utf-8") as f:
                    self.config = json.load(f)
            except Exception:
                self.config = {"show_instruction": True}

    def save_config(self):
        try:
            with open(CONFIG_FILE, "w", encoding="utf-8") as f:
                json.dump(self.config, f)
        except Exception:
            pass

    def show_instruction_if_needed(self):
        # 不再判断show_instruction，每次都弹窗
        self.show_instruction_dialog()

    def show_instruction_dialog(self):
        win = tk.Toplevel(self.root)
        win.title("使用说明")
        win.geometry("480x400")
        win.grab_set()
        win.transient(self.root)
        win.protocol("WM_DELETE_WINDOW", lambda: None)  # 禁用关闭按钮
        text = (
            "使用说明：\n"
            "1. 先点击'选择文档目录'选择需要合成的文档所在文件夹。\n"
            "2. 再点击'命名合并文档'输入合成后文档的文件名。\n"
            "3. 可选择保存的目标目录（不选则默认保存到文档目录）。\n"
            "4. 最后点击'开始合成'即可。\n"
            "5. 合成时会自动分页，保留所有内容（包括图片、表格、格式等）。\n"
            "6. 仅支持.doc和.docx文件。\n"
            "7. 所有需要合成的文档必须全部关闭，否则可能导致合成失败。\n"
            "8. 如遇异常请联系管理员。"
        )
        label = tk.Label(win, text=text, justify="left", anchor="nw", font=LABEL_FONT, bg=BG_COLOR, fg=LABEL_COLOR)
        label.pack(padx=24, pady=24, fill="both", expand=True)

        var_no_more = tk.BooleanVar(value=False)
        var_next_time = tk.BooleanVar(value=False)

        def on_check():
            if var_no_more.get() or var_next_time.get():
                btn.config(state="normal")
            else:
                btn.config(state="disabled")

        def on_confirm():
            self.save_config()
            win.destroy()

        cb1 = tk.Checkbutton(win, text="知道啦，不用再说了", variable=var_no_more, command=lambda: [var_next_time.set(False), on_check()], font=LABEL_FONT, bg=BG_COLOR)
        cb2 = tk.Checkbutton(win, text="下次使用我还要看看呢", variable=var_next_time, command=lambda: [var_no_more.set(False), on_check()], font=LABEL_FONT, bg=BG_COLOR)
        cb1.pack(anchor="w", padx=36)
        cb2.pack(anchor="w", padx=36)
        btn = tk.Button(win, text="确定", command=on_confirm, font=BTN_FONT, bg=PRIMARY_COLOR, fg="#fff", activebackground=ACCENT_COLOR, activeforeground="#fff", state="disabled")
        btn.pack(pady=16)

    def safe_call(self, func):
        def wrapper(*args, **kwargs):
            try:
                return func(*args, **kwargs)
            except Exception as e:
                messagebox.showerror("操作异常", f"操作异常，请按要求操作！\n\n详细信息：{str(e)}")
        return wrapper

    def select_directory(self):
        directory = filedialog.askdirectory()
        if not directory:
            return
        self.selected_dir = directory
        self.dir_label.config(text=f"已选择目录: {directory}")
        self.name_btn.config(state="normal")
        self.save_dir_btn.config(state="normal")
        self.merge_btn.config(state="disabled")
        self.output_name = None
        self.save_dir = None
        self.status_label.config(text="")

    def set_output_name(self):
        if not self.selected_dir:
            messagebox.showerror("错误", "请先选择文档目录！")
            return
        output_name = simpledialog.askstring("输入", "请输入合并后的文件名（不需要扩展名）：")
        if output_name:
            self.output_name = output_name
            self.merge_btn.config(state="normal")
            self.status_label.config(text=f"已设置文件名: {output_name}.docx")
        else:
            self.merge_btn.config(state="disabled")
            self.status_label.config(text="")

    def select_save_directory(self):
        if not self.selected_dir:
            messagebox.showerror("错误", "请先选择文档目录！")
            return
        save_dir = filedialog.askdirectory(title="选择保存的目标目录")
        if save_dir:
            self.save_dir = save_dir
            self.status_label.config(text=f"已选择保存目录: {save_dir}")
        else:
            self.save_dir = None
            self.status_label.config(text="未选择保存目录，默认保存到文档目录")

    def append_docx_content(self, src_doc, dest_doc):
        # 复制src_doc的body内容到dest_doc，保留图片、表格、格式等
        for element in src_doc.element.body:
            dest_doc.element.body.append(element)

    def merge_documents(self):
        if not self.selected_dir:
            messagebox.showerror("错误", "请先选择文档目录！")
            return
        if not self.output_name:
            messagebox.showerror("错误", "请先命名合并文档！")
            return
        output_dir = self.save_dir if self.save_dir else self.selected_dir
        # 自动创建目标文件夹
        if not os.path.exists(output_dir):
            try:
                os.makedirs(output_dir)
            except Exception as e:
                messagebox.showerror("错误", f"无法创建目标文件夹：\n{output_dir}\n{str(e)}")
                return
        output_path = os.path.join(output_dir, f"{self.output_name}.docx")
        doc_files = []
        for file in os.listdir(self.selected_dir):
            if any(file.lower().endswith(ext) for ext in self.supported_extensions):
                doc_files.append(os.path.join(self.selected_dir, file))
        if not doc_files:
            messagebox.showerror("错误", "所选目录中没有找到支持的文档文件！")
            return
        try:
            merged_doc = Document(doc_files[0])  # 以第一个文档为模板
            total_files = len(doc_files)
            self.progress["value"] = 0
            self.status_label.config(text="开始合成...", style="Status.TLabel")
            self.root.update()
            # 清空模板文档内容，只保留结构
            merged_doc._body.clear_content()
            for idx, doc_file in enumerate(doc_files):
                self.progress["value"] = (idx + 1) / total_files * 100
                self.status_label.config(text=f"正在处理: {os.path.basename(doc_file)}", style="Status.TLabel")
                self.root.update()
                if doc_file.endswith('.docx'):
                    doc = Document(doc_file)
                else:
                    doc = self.convert_doc_to_docx(doc_file)
                self.append_docx_content(doc, merged_doc)
                # 只在文档之间插入分页符，最后一个文档后不插
                if idx < total_files - 1:
                    merged_doc.add_page_break()
            merged_doc.save(output_path)
            self.status_label.config(text="文档合成完成！", style="Success.TLabel")
            messagebox.showinfo("成功", f"文档已成功合成到：\n{output_path}")
        except Exception as e:
            self.status_label.config(text="合成失败！", style="Error.TLabel")
            messagebox.showerror("错误", f"合成过程中出现错误：\n{str(e)}")
        finally:
            self.progress["value"] = 0
            self.status_label.config(text="")

    def convert_doc_to_docx(self, doc_path):
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        try:
            doc = word.Documents.Open(doc_path)
            temp_docx = doc_path + "x"
            doc.SaveAs2(temp_docx, FileFormat=16)
            doc.Close()
            return Document(temp_docx)
        finally:
            word.Quit()
            if os.path.exists(temp_docx):
                os.remove(temp_docx)

    def reset_all(self):
        self.selected_dir = None
        self.output_name = None
        self.save_dir = None
        self.dir_label.config(text="未选择目录")
        self.status_label.config(text="", style="Status.TLabel")
        self.name_btn.config(state="disabled")
        self.save_dir_btn.config(state="disabled")
        self.merge_btn.config(state="disabled")
        self.progress["value"] = 0

    def on_close(self):
        self.save_config()
        self.root.destroy()

def main():
    try:
        root = tk.Tk()
        app = DocumentMerger(root)
        root.mainloop()
    except Exception as e:
        messagebox.showerror("错误", f"程序运行出错：\n{str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main() 