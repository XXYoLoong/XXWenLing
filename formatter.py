import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import re
from docx.oxml.ns import qn
from database.models import DocTask, TaskLog, PerformanceLog, FormatTemplate
from datetime import datetime

class DocumentFormatter:
    def __init__(self, user_id=None):
        self.user_id = user_id
        self.current_template = None
        
    def set_current_template(self, name):
        ft = FormatTemplate()
        templates = ft.get_user_templates(self.user_id)
        for t in templates:
            if t['name'] == name:
                self.current_template = json.loads(t['config'])
                return True
        return False
        
    def classify_paragraph(self, paragraph, rules, para_index=1):
        """根据规则对段落进行分类，para_index为1基序号"""
        text = paragraph.text.strip()
        for rule_type, rule in rules.items():
            if rule['type'] == 'regex' and re.search(rule['pattern'], text):
                return rule_type
            elif rule['type'] == 'keyword' and rule['keyword'] in text:
                return rule_type
            elif rule['type'] == 'position' and rule['position'] == para_index:
                return rule_type
            elif rule['type'] == 'length' and len(text) <= rule['max_length']:
                return rule_type
        return 'body'  # 默认为正文

    def clear_format(self, paragraph):
        # 清空段落和run的所有格式
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.left_indent = None
        paragraph.paragraph_format.right_indent = None
        paragraph.paragraph_format.space_before = None
        paragraph.paragraph_format.space_after = None
        paragraph.paragraph_format.line_spacing = None
        paragraph.alignment = None
        for run in paragraph.runs:
            run.font.name = None
            run.font.size = None
            run.font.bold = None
            run.font.italic = None
            run.font.underline = None

    def apply_format(self, paragraph, format_rules):
        self.clear_format(paragraph)
        if 'font' in format_rules:
            for run in paragraph.runs:
                run.font.name = format_rules['font']
                try:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), format_rules['font'])
                except Exception:
                    pass
        if 'size' in format_rules:
            for run in paragraph.runs:
                run.font.size = Pt(format_rules['size'])
        if 'alignment' in format_rules:
            alignment_map = {
                '左对齐': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                '居中': WD_ALIGN_PARAGRAPH.CENTER,
                '右对齐': WD_ALIGN_PARAGRAPH.RIGHT,
                '两端对齐': WD_ALIGN_PARAGRAPH.JUSTIFY,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            paragraph.alignment = alignment_map.get(format_rules['alignment'], WD_ALIGN_PARAGRAPH.LEFT)
        if 'indent' in format_rules:
            paragraph.paragraph_format.first_line_indent = Cm(format_rules['indent'])
        if 'line_spacing' in format_rules:
            paragraph.paragraph_format.line_spacing = Pt(format_rules['line_spacing'])
        # 字体样式
        if 'bold' in format_rules:
            for run in paragraph.runs:
                run.font.bold = format_rules['bold']
        if 'italic' in format_rules:
            for run in paragraph.runs:
                run.font.italic = format_rules['italic']
        if 'underline' in format_rules:
            for run in paragraph.runs:
                run.font.underline = format_rules['underline']

    def format_document(self, doc_path, output_path=None, progress_callback=None, status_callback=None):
        """格式化单个文档"""
        try:
            doc = Document(doc_path)
            template = self.current_template
            
            if not template:
                raise Exception("未选择格式模板！")
                
            total_paragraphs = len(doc.paragraphs)
            
            for i, paragraph in enumerate(doc.paragraphs):
                if progress_callback:
                    progress_callback(int((i + 1) / total_paragraphs * 100))
                if status_callback:
                    status_callback(f"正在处理第 {i+1}/{total_paragraphs} 段")
                    
                # 不跳过空段落，位置规则严格与文档段落索引对应
                para_type = self.classify_paragraph(paragraph, template['rules'], para_index=i+1)
                
                # 只对非空段落应用格式
                if para_type in template['formats'] and paragraph.text.strip():
                    self.apply_format(paragraph, template['formats'][para_type])
                    
            # 保存文档
            if output_path:
                doc.save(output_path)
            else:
                base, ext = os.path.splitext(doc_path)
                doc.save(f"{base}_formatted{ext}")
                
            return True
            
        except Exception as e:
            raise Exception(f"处理文档时出错: {str(e)}")
            
    def format_directory(self, dir_path, output_dir=None, progress_callback=None, status_callback=None):
        """批量处理目录下的所有文档"""
        if not os.path.exists(dir_path):
            raise Exception("目录不存在！")
            
        doc_files = [f for f in os.listdir(dir_path) if f.endswith('.docx')]
        if not doc_files:
            raise Exception("目录中没有.docx文件！")
            
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
            
        total_files = len(doc_files)
        success_count = 0
        failed_files = []
        
        for i, doc_file in enumerate(doc_files):
            try:
                doc_path = os.path.join(dir_path, doc_file)
                if output_dir:
                    output_path = os.path.join(output_dir, f"{os.path.splitext(doc_file)[0]}_formatted.docx")
                else:
                    output_path = None
                    
                if progress_callback:
                    progress_callback(int((i + 1) / total_files * 100))
                if status_callback:
                    status_callback(f"正在处理: {doc_file}")
                    
                self.format_document(doc_path, output_path)
                success_count += 1
                
            except Exception as e:
                failed_files.append((doc_file, str(e)))
                
        return {
            'total': total_files,
            'success': success_count,
            'failed': failed_files
        }

    def format_files(self, file_list, output_dir=None, progress_callback=None, status_callback=None):
        total_files = len(file_list)
        success_count = 0
        failed_files = []
        task_model = DocTask()
        task_name = f"批量格式化_{datetime.now().strftime('%Y%m%d%H%M%S')}"
        output_path = output_dir or ''
        task_id = task_model.create(self.user_id, 'format', task_name, '', output_path, None)
        log_model = TaskLog()
        perf_model = PerformanceLog()
        start_time = datetime.now()
        for i, doc_path in enumerate(file_list):
            try:
                if output_dir:
                    output_path = os.path.join(output_dir, f"{os.path.splitext(os.path.basename(doc_path))[0]}_formatted.docx")
                else:
                    output_path = None
                if progress_callback:
                    progress_callback(int((i + 1) / total_files * 100))
                if status_callback:
                    status_callback(f"正在处理: {os.path.basename(doc_path)}")
                log_model.add_log(task_id, 'progress', f"正在处理: {os.path.basename(doc_path)}", user_id=self.user_id)
                self.format_document(doc_path, output_path)
                success_count += 1
            except Exception as e:
                failed_files.append((os.path.basename(doc_path), str(e)))
                log_model.add_log(task_id, 'error', f"处理失败: {os.path.basename(doc_path)} - {str(e)}", user_id=self.user_id)
        end_time = datetime.now()
        duration = int((end_time - start_time).total_seconds() * 1000)
        perf_model.add_log(task_id, 'format', duration, None, user_id=self.user_id)
        if failed_files:
            task_model.update_status(task_id, 'failed', f"失败文件数: {len(failed_files)}", user_id=self.user_id)
        else:
            task_model.update_status(task_id, 'success', "全部处理成功", user_id=self.user_id)
        log_model.add_log(task_id, 'info', f"处理完成，成功: {success_count}，失败: {len(failed_files)}", user_id=self.user_id)
        return {
            'total': total_files,
            'success': success_count,
            'failed': failed_files
        } 