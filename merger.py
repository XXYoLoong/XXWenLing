import os
from docx import Document
import pythoncom
import win32com.client
from database.models import DocTask, TaskLog, PerformanceLog
from datetime import datetime

def merge_documents(doc_dir, out_name, save_dir, progress=None, status=None, user_id=None):
    supported_extensions = ['.doc', '.docx']
    doc_files = [os.path.join(doc_dir, f) for f in os.listdir(doc_dir) if any(f.lower().endswith(ext) for ext in supported_extensions)]
    if not doc_files:
        raise Exception("所选目录中没有支持的文档文件！")
    if not os.path.exists(save_dir):
        os.makedirs(save_dir)
    output_path = os.path.join(save_dir, f"{out_name}.docx")
    merged_doc = Document(doc_files[0])
    merged_doc._body.clear_content()
    total = len(doc_files)
    # 创建数据库任务记录
    task_model = DocTask()
    task_id = task_model.create(user_id, 'merge', out_name, doc_dir, output_path, None)
    log_model = TaskLog()
    perf_model = PerformanceLog()
    start_time = datetime.now()
    try:
        for idx, doc_file in enumerate(doc_files):
            if progress:
                progress.setValue(int((idx + 1) / total * 100))
            if status:
                status.setText(f"正在处理: {os.path.basename(doc_file)}")
            log_model.add_log(task_id, 'progress', f"正在处理: {os.path.basename(doc_file)}")
            if doc_file.endswith('.docx'):
                doc = Document(doc_file)
            else:
                doc = convert_doc_to_docx(doc_file)
            for element in doc.element.body:
                merged_doc.element.body.append(element)
            if idx < total - 1:
                merged_doc.add_page_break()
        merged_doc.save(output_path)
        end_time = datetime.now()
        duration = int((end_time - start_time).total_seconds() * 1000)
        perf_model.add_log(task_id, 'merge', duration, None)
        task_model.update_status(task_id, 'success', f"合成完成：{output_path}")
        log_model.add_log(task_id, 'info', f"合成完成：{output_path}")
        return output_path
    except Exception as e:
        task_model.update_status(task_id, 'failed', str(e))
        log_model.add_log(task_id, 'error', f"合成失败：{str(e)}")
        raise

def convert_doc_to_docx(doc_path):
    pythoncom.CoInitialize()
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    temp_docx = doc_path + "x"
    try:
        doc = word.Documents.Open(doc_path)
        doc.SaveAs2(temp_docx, FileFormat=16)
        doc.Close()
        return Document(temp_docx)
    finally:
        word.Quit()
        if os.path.exists(temp_docx):
            os.remove(temp_docx) 